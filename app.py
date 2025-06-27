# app.py
# -*- coding: utf-8 -*-

# =============================================================================
# SECTION 0: IMPORTS
# =============================================================================
import os
import traceback
import random
import re
import docx  
import fitz  
import time
import json
from datetime import datetime, date, timedelta, timezone
from flask import Flask, render_template, request, redirect, url_for, flash, send_from_directory, abort, jsonify
from flask_sqlalchemy import SQLAlchemy
from flask_migrate import Migrate
from werkzeug.utils import secure_filename
from sqlalchemy import or_, and_, desc
from sqlalchemy import func as sql_func
from sqlalchemy.exc import SQLAlchemyError
from werkzeug.security import generate_password_hash

# =============================================================================
# SECTION 1: FLASK APP INITIALIZATION & CONFIGURATION
# =============================================================================
app = Flask(__name__)
basedir = os.path.abspath(os.path.dirname(__file__))

# --- Basic Configs ---
app.config['UPLOAD_FOLDER'] = os.path.join(basedir, 'uploads/')
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', 'a_default_secret_key_for_development')
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///' + os.path.join(basedir, 'studyvault.db')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# --- File Type Configs ---
app.config['ALLOWED_EXTENSIONS'] = {'txt', 'pdf', 'docx'}
app.config['ALLOWED_IMAGE_EXTENSIONS'] = {'png', 'jpg', 'jpeg', 'gif', 'webp'}
app.config['ALLOWED_VIDEO_EXTENSIONS'] = {'mp4', 'mov', 'avi', 'mkv'}

# --- App Initialization ---
db = SQLAlchemy(app)
migrate = Migrate(app, db)
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# --- Global Constants ---
ITEMS_PER_PAGE = 10
STOP_WORDS = set(["là", "và", "của", "có", "trong", "để", "một", "không", "được", "cho", "với", "tại","thì", "mà", "khi", "từ", "ra", "lên", "xuống", "vào", "qua", "đến", "đi", "lại","như", "ở", "đã", "sẽ", "đang", "rằng", "hay", "hơn", "rất", "này", "đó", "kia", "ấy","tôi", "bạn", "anh", "chị", "em", "ông", "bà", "nó", "chúng", "mình","the", "a", "an", "is", "are", "was", "were", "of", "in", "on", "at", "to", "for","with", "by", "from", "as", "and", "or", "but", "if", "then", "this", "that", "it","its", "i", "you", "he", "she", "we", "they", "my", "your", "his", "her", "our", "their"])

# =============================================================================
# SECTION 2: DATABASE MODELS
# =============================================================================
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    email = db.Column(db.String(120), unique=True, nullable=False)
    password_hash = db.Column(db.String(128), nullable=False)
    ultimate_goal = db.Column(db.Text, nullable=True)
    role_model_character = db.Column(db.String(100), nullable=True)
    workspace_vibe = db.Column(db.String(100), nullable=True, default='Study')
    workspace_color_theme = db.Column(db.String(50), nullable=True, default='blue')
    selected_avatar = db.Column(db.String(100), nullable=True)
    short_term_mode_active = db.Column(db.Boolean, default=False)
    short_term_mode_end_date = db.Column(db.Date, nullable=True)
    short_term_mode_focus_keywords = db.Column(db.Text, nullable=True)
    short_term_mode_intensity = db.Column(db.String(10), nullable=True)
    short_term_mode_last_used = db.Column(db.Date, nullable=True)
    specific_study_goal = db.Column(db.Text, nullable=True)
    preferred_content_types = db.Column(db.Text, nullable=True)
    expected_completion_time = db.Column(db.String(50), nullable=True)
    personal_learning_challenges = db.Column(db.Text, nullable=True)
    studyvault_expectations = db.Column(db.Text, nullable=True)

    def __repr__(self):
        return f'<User {self.username}>'

class Document(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(200), nullable=False)
    filepath = db.Column(db.String(300), nullable=False, unique=True)
    category = db.Column(db.String(100), nullable=True, default="Chưa phân loại")
    uploaded_date = db.Column(db.DateTime, nullable=False, default=lambda: datetime.now(timezone.utc))
    last_viewed_date = db.Column(db.DateTime, nullable=True)
    doc_type = db.Column(db.String(20), nullable=False, default='file')
    engagement_level = db.Column(db.String(50), nullable=True)
    custom_note = db.Column(db.Text, nullable=True)
    deadline = db.Column(db.Date, nullable=True)
    extracted_content = db.Column(db.Text, nullable=True)
    user_summary = db.Column(db.Text, nullable=True)
    ai_topic_label = db.Column(db.String(100), nullable=True)
    win_criteria_description = db.Column(db.Text, nullable=True)
    target_score = db.Column(db.Integer, nullable=True)
    actual_score = db.Column(db.Integer, nullable=True)
    keywords = db.Column(db.Text, nullable=True)
    context_event = db.Column(db.String(200), nullable=True)
    is_goal_related = db.Column(db.Boolean, default=False, nullable=False)
    filename_normalized = db.Column(db.String(200), nullable=True)
    workspace_items = db.relationship('WorkspaceItem', backref='document', lazy=True, cascade="all, delete-orphan")
    relations = db.relationship('WorkspaceItemRelation', backref='document', lazy=True, cascade="all, delete-orphan")
    learning_objectives = db.relationship('LearningObjective', backref='doc', lazy=True, cascade="all, delete-orphan")

class WorkspaceItem(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    title = db.Column(db.String(300), nullable=False)
    content = db.Column(db.Text, nullable=True)
    order = db.Column(db.Integer, nullable=False, default=0)
    user_content = db.Column(db.Text, nullable=True)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False)
    parent_id = db.Column(db.Integer, db.ForeignKey('workspace_item.id'), nullable=True)
    children = db.relationship('WorkspaceItem', backref=db.backref('parent', remote_side=[id]), lazy='dynamic', cascade="all, delete-orphan")

class WorkspaceItemRelation(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=False)
    source_id = db.Column(db.Integer, db.ForeignKey('workspace_item.id'), nullable=False)
    target_id = db.Column(db.Integer, db.ForeignKey('workspace_item.id'), nullable=False)
    label = db.Column(db.String(100), nullable=True)
    source_node = db.relationship('WorkspaceItem', foreign_keys=[source_id], backref='source_relations')
    target_node = db.relationship('WorkspaceItem', foreign_keys=[target_id], backref='target_relations')

class LearningObjective(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    description = db.Column(db.String(500), nullable=False)
    is_completed = db.Column(db.Boolean, default=False, nullable=False)
    document_id = db.Column(db.Integer, db.ForeignKey('document.id'), nullable=True)
    parent_id = db.Column(db.Integer, db.ForeignKey('learning_objective.id'), nullable=True)
    sub_objectives = db.relationship('LearningObjective', backref=db.backref('parent', remote_side=[id]), lazy=True, cascade="all, delete-orphan")
    importance = db.Column(db.String(50), nullable=True)
    learning_role = db.Column(db.String(50), nullable=True)
    difficulty = db.Column(db.String(50), nullable=True)

    def __repr__(self):
        return f'<LearningObjective {self.id}: {self.description[:30]}>'
    
# =============================================================================
# SECTION 3: HELPER CLASSES & UTILITY FUNCTIONS
# =============================================================================
class FileProcessor:
    DEFAULT_CATEGORY = "Chưa phân loại"
    CATEGORY_KEYWORDS = {
        "Lập trình": ["python", "java", "code", "script", "lập trình", "thuật toán", "hàm", "biến", "vòng lặp"],
        "Kinh tế học": ["kinh tế", "thị trường", "gdp", "lạm phát", "cung", "cầu", "vi mô", "vĩ mô"],
        "Toán học": ["toán", "công thức", "phương trình", "tích phân", "đạo hàm", "ma trận", "vector"],
        "Ngoại ngữ": ["english", "tiếng anh", "từ vựng", "grammar", "ielts", "toeic"],
        "Kỹ năng mềm": ["cv", "giao tiếp", "thuyết trình", "lãnh đạo"],
        "Vật Lý": ["điện trường", "điện dung"],
        "Ngữ Văn": ["tác phẩm", "thơ", "truyện"],
        "Hóa học": ["nguyên tố", "hóa chất", "phân tử khối"],
        "Sinh học": ["tế bào", "sinh vật", "động vật", "thực vậtvật"],
        "Pháp luật": ["luật", "hiến pháp", "điều khoản", "nghị định", "thông tư"],
        "AI/ML": ["machine learning", "ai", "neural network", "mạng nơ-ron", "học máy"],
        "Link Web": ["http", "https"],
        "Hình ảnh": ["image", "picture"],
        "Video": ["video", "media"],
        "Google Drive": ["google drive", "gsheet"]
    }

    def extract_text(self, filepath):
        try:
            ext = os.path.splitext(filepath)[1].lower()
            if ext == '.pdf':
                with fitz.open(filepath) as doc:
                    text = "".join(page.get_text() for page in doc)
                return text
            elif ext == '.docx':
                doc = docx.Document(filepath)
                return "\n".join([para.text for para in doc.paragraphs])
            elif ext == '.txt':
                with open(filepath, 'r', encoding='utf-8') as f:
                    return f.read()
            else:
                return f"[Lỗi: Định dạng file '{ext}' không được hỗ trợ để trích xuất văn bản.]"
        except Exception as e:
            print(f"ERROR extracting text from {filepath}: {e}")
            return f"[Lỗi: Không thể đọc nội dung từ file. File có thể bị hỏng hoặc cần quyền truy cập. Chi tiết: {e}]"

    def categorize_document(self, filename_or_content, user_ultimate_goal=None, user_role_model=None):
        keyword_string = (filename_or_content or '').lower()  
        detected_categories = set()
        for cat, keys in self.CATEGORY_KEYWORDS.items():
            if any(k in keyword_string for k in keys):
                detected_categories.add(cat)
        if user_ultimate_goal:
            normalized_goal = normalize_vietnamese(user_ultimate_goal).lower()
            if "ai" in normalized_goal or "hoc may" in normalized_goal or "lap trinh" in normalized_goal:
                if "neural network" in keyword_string or "machine learning" in keyword_string:
                    detected_categories.add("AI/ML")
                if "python" in keyword_string or "code" in keyword_string:
                    detected_categories.add("Lập trình")
            if "kinh doanh" in normalized_goal or "doanh nhan" in normalized_goal:
                if "thi truong" in keyword_string or "kinh te" in keyword_string:
                    detected_categories.add("Kinh tế học")
        
        if user_role_model:
            normalized_role = normalize_vietnamese(user_role_model).lower()
            if "chuyen gia ai" in normalized_role or "game developer" in normalized_role:
                if "thuật toán" in keyword_string or "code" in keyword_string:
                    detected_categories.add("Lập trình")
            if "doanh nhan thanh dat" in normalized_role:
                if "gdp" in keyword_string or "thị trường" in keyword_string:
                    detected_categories.add("Kinh tế học")
        
        if detected_categories:
            if "AI/ML" in detected_categories: return "AI/ML"
            if "Lập trình" in detected_categories: return "Lập trình"
            if "Kinh tế học" in detected_categories: return "Kinh tế học"
            if "Toán học" in detected_categories: return "Toán học"
            return sorted(list(detected_categories))[0] 
        
        return self.DEFAULT_CATEGORY

    @staticmethod
    def build_objectives_tree(items): 
        tree = []
        item_map = {item.id: {
            "id": item.id,
            "description": item.description,
            "is_completed": item.is_completed,
            "sub_objectives": []
        } for item in items}
        for item in items:
            if item.parent_id is not None and item.parent_id in item_map: 
                parent_item_data = item_map[item.parent_id]
                parent_item_data["sub_objectives"].append(item_map[item.id])
            else:   
                tree.append(item_map[item.id])
        return tree
    
# --- Khởi tạo đối tượng xử lý file ---
fp = FileProcessor()

# --- Các hàm tiện ích (Utility Functions) ---
def normalize_vietnamese(text):
    if not text: return ""
    replacements = {'á': 'a', 'à': 'a', 'ả': 'a', 'ã': 'a', 'ạ': 'a','ă': 'a', 'ằ': 'a', 'ắ': 'a', 'ẳ': 'a', 'ẵ': 'a', 'ặ': 'a','â': 'a', 'ầ': 'a', 'ấ': 'a', 'ẩ': 'a', 'ẫ': 'a', 'ậ': 'a','đ': 'd','é': 'e', 'è': 'e', 'ẻ': 'e', 'ẽ': 'e', 'ẹ': 'e','ê': 'e', 'ề': 'e', 'ế': 'e', 'ể': 'e', 'ễ': 'e', 'ệ': 'e','í': 'i', 'ì': 'i', 'ỉ': 'i', 'ĩ': 'i', 'ị': 'i','ó': 'o', 'ò': 'o', 'ỏ': 'o', 'õ': 'o', 'ọ': 'o','ô': 'o', 'ồ': 'o', 'ố': 'o', 'ổ': 'o', 'ỗ': 'o', 'ộ': 'o','ơ': 'o', 'ờ': 'o', 'ớ': 'o', 'ở': 'o', 'ỡ': 'o', 'ợ': 'o','ú': 'u', 'ù': 'u', 'ủ': 'u', 'ũ': 'u', 'ụ': 'u','ư': 'u', 'ừ': 'u', 'ứ': 'u', 'ử': 'u', 'ữ': 'u', 'ự': 'u','ý': 'y', 'ỳ': 'y', 'ỷ': 'y', 'ĩ': 'y', 'ỵ': 'y',}
    text_lower = text.lower()
    return "".join([replacements.get(char, char) for char in text_lower])

def allowed_file(filename): return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']
def allowed_image(filename): return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_IMAGE_EXTENSIONS']
def allowed_video(filename): return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_VIDEO_EXTENSIONS']
def is_google_drive_link(url): return url and "drive.google.com" in url.lower()

def get_random_docs(model, num_docs):
    all_ids = db.session.query(model.id).all()
    if not all_ids: return []
    random_id_tuples = random.sample(all_ids, min(len(all_ids), num_docs))
    random_ids = [r[0] for r in random_id_tuples]
    return model.query.filter(model.id.in_(random_ids)).all()

def get_unique_random_elements(input_list, num_elements):
    if not input_list: return []
    return random.sample(input_list, min(len(input_list), num_elements))
STOP_WORDS = set(["là", "và", "của", "có", "trong", "để", "một", "không", "được", "cho", "với", "tại","thì", "mà", "khi", "từ", "ra", "lên", "xuống", "vào", "qua", "đến", "đi", "lại","như", "ở", "đã", "sẽ", "đang", "rằng", "hay", "hơn", "rất", "này", "đó", "kia", "ấy","tôi", "bạn", "anh", "chị", "em", "ông", "bà", "nó", "chúng", "mình","the", "a", "an", "is", "are", "was", "were", "of", "in", "on", "at", "to", "for","with", "by", "from", "as", "and", "or", "but", "if", "then", "this", "that", "it","its", "i", "you", "he", "she", "we", "they", "my", "your", "his", "her", "our", "their"])

def create_fill_in_the_blank_question(content_text, num_blanks=1, min_word_len=4):
    if not content_text or not isinstance(content_text, str): return None
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?|\!)\s', content_text)
    potential_sentences = [s.strip() for s in sentences if len(s.strip().split()) > 5]
    if not potential_sentences: return None
    random.shuffle(potential_sentences)
    for sentence in potential_sentences:
        words = re.findall(r'\b\w+\b', sentence.lower())
        candidate_keywords = [word for word in words if len(word) >= min_word_len and word not in STOP_WORDS]
        if len(candidate_keywords) < num_blanks: continue
        words_to_blank = random.sample(candidate_keywords, num_blanks)
        original_sentence = sentence; question_sentence = sentence
        original_words_for_answer = []
        words_to_blank.sort(key=len, reverse=True)
        for word_to_blank_lower in words_to_blank:
            match = re.search(r'\b' + re.escape(word_to_blank_lower) + r'\b', question_sentence, re.IGNORECASE)
            if match:
                original_word = match.group(0)
                new_question_sentence = re.sub(r'\b' + re.escape(original_word) + r'\b', "______", question_sentence, 1, flags=re.IGNORECASE)
                if new_question_sentence != question_sentence:
                    question_sentence = new_question_sentence
                    original_words_for_answer.append(original_word)
        if len(original_words_for_answer) == num_blanks:
            return {"q": question_sentence, "a": " / ".join(original_words_for_answer), "original_sentence": original_sentence}
    return None

def build_tree(items):
    tree = []
    item_map = {item.id: {
        "id": item.id,
        "title": item.title,
        "content": item.content,
        "order": item.order,
        "user_content": item.user_content,
        "children": []
    } for item in items}
    for item in items:
        if item.parent_id is not None and item.parent_id in item_map:  
            parent_item = item_map[item.parent_id]
            parent_item["children"].append(item_map[item.id])
        else: 
            tree.append(item_map[item.id])
    for item_data in item_map.values():
        item_data['children'].sort(key=lambda x: x['order'])
    tree.sort(key=lambda x: x['order'])
    return tree
   
def generate_mermaid_graph(items):
    if not items:
        return "Chưa có mục nào trong không gian làm việc. Hãy tạo mục đầu tiên!"
    mermaid_string = "graph TD;\n"
    relationships = [
        "-->|dẫn tới|", 
        "-.->|so sánh|", 
        "==>|kết quả|", 
        "--o|là một phần của|"
    ]
    colors = ["#87CEEB", "#FFD700", "#98FB98", "#FFB6C1", "#F0E68C"]
    for item in items:
        clean_title = item.title.replace('"', '#quot;').replace('\n', '<br/>')
        mermaid_string += f'    node_{item.id}["{clean_title}"];\n'
    link_index = 0
    link_styles = ""
    for item in items:
        if item.parent_id:
            random_relationship = random.choice(relationships)
            mermaid_string += f'    node_{item.parent_id} {random_relationship} node_{item.id};\n'
            random_color = random.choice(colors)
            link_styles += f'    linkStyle {link_index} stroke:{random_color},stroke-width:2px;\n'
            link_index += 1  
    mermaid_string += "\n" + link_styles
    return mermaid_string

def check_document_relevance(document, user):
    """
    Kiểm tra xem một tài liệu có liên quan đến mục tiêu của người dùng không.
    Trả về True nếu liên quan, False nếu không.
    """
    if not user or not document:
        return False

    user_keywords = set()
    if user.ultimate_goal or user.role_model_character or user.specific_study_goal:
        goal_text = (user.ultimate_goal or '') + ' ' + (user.role_model_character or '') + ' ' + (user.specific_study_goal or '')
        normalized_goal_text = normalize_vietnamese(goal_text).lower()
        words = normalized_goal_text.split()
        user_keywords = {word for word in words if word not in STOP_WORDS and len(word) >= 2}

    if not user_keywords:
        return False

    doc_text_for_check = (
        (document.filename_normalized or normalize_vietnamese(document.filename)) + ' ' +
        (normalize_vietnamese(document.category) if document.category else '') + ' ' +
        (normalize_vietnamese(document.keywords) if document.keywords else '')
    ).lower()

    if any(re.search(r'\b' + re.escape(keyword) + r'\b', doc_text_for_check) for keyword in user_keywords):
        return True
    
    return False

# =============================================================================
# SECTION 4: FLASK ROUTES
# =============================================================================

#---layout---
@app.context_processor
def inject_user():
    try:
        user = User.query.first()
        return dict(user=user)
    except Exception as e:
        print(f"Error injecting user: {e}")
        return dict(user=None)

# --- Main Page Routes ---
@app.route('/')
def index():
    current_user = User.query.first()
    pagination_data = None
    documents_on_page = []
    suggested_docs = []
    documents_data_for_js = [] 

    search_query = request.args.get('search_query', '').strip()
    category_filter = request.args.get('category', '').strip()
    available_categories = sorted(list(fp.CATEGORY_KEYWORDS.keys()))
    if fp.DEFAULT_CATEGORY not in available_categories:
        available_categories.append(fp.DEFAULT_CATEGORY)
        available_categories.sort()

    if not current_user:
        try:
            new_user = User(username='demo_user', email='demo@example.com', password_hash=generate_password_hash('password'))
            db.session.add(new_user)
            db.session.commit()
            current_user = new_user
            flash("Chào mừng đến với StudyVault! Vui lòng thiết lập hồ sơ của bạn.", "info")
        except Exception as e:
            print(f"Lỗi khi tạo người dùng demo trong index: {e}")
            
    try:
        page = request.args.get('page', 1, type=int)
        query = Document.query
        if search_query:
            normalized_search_query = normalize_vietnamese(search_query)
            query = query.filter(or_(Document.filename.ilike(f"%{search_query}%"), Document.filename_normalized.ilike(f"%{normalized_search_query}%"), Document.keywords.ilike(f"%{search_query}%")))
        if category_filter:
            query = query.filter(Document.category == category_filter)
        query = query.order_by(Document.uploaded_date.desc())
        pagination_data = query.paginate(page=page, per_page=ITEMS_PER_PAGE, error_out=False)
        documents_on_page = pagination_data.items if pagination_data else []
        
        # Giả lập việc xác định tài liệu liên quan đến mục tiêu
        for doc in documents_on_page:
            doc.is_goal_related = check_document_relevance(doc, current_user)

        # Phần gợi ý tài liệu
        try:
            base_query_for_suggestions = Document.query.filter(
                or_(Document.last_viewed_date.is_(None), Document.engagement_level.is_(None), Document.context_event.is_(None))
            )
            if current_user and (current_user.ultimate_goal or current_user.role_model_character):
                goal_keywords_for_suggestions = []
                if current_user.ultimate_goal:
                    goal_keywords_for_suggestions.extend(normalize_vietnamese(current_user.ultimate_goal).lower().split())
                if current_user.role_model_character:
                    goal_keywords_for_suggestions.extend(normalize_vietnamese(current_user.role_model_character).lower().split())
                
                goal_conditions_for_suggestions = []
                for keyword in set(goal_keywords_for_suggestions):
                    goal_conditions_for_suggestions.append(or_(
                        Document.filename_normalized.ilike(f"%{keyword}%"),
                        Document.keywords.ilike(f"%{keyword}%"),
                        Document.category.ilike(f"%{keyword}%")
                    ))
                
                if goal_conditions_for_suggestions:
                    goal_related_docs = base_query_for_suggestions.filter(or_(*goal_conditions_for_suggestions)).limit(3).all()
                    suggested_docs.extend(goal_related_docs)
                    
                    needed_more = 3 - len(suggested_docs)
                    if needed_more > 0:
                        existing_ids = [doc.id for doc in suggested_docs]
                        random_docs = base_query_for_suggestions.filter(Document.id.notin_(existing_ids)).limit(needed_more).all()
                        suggested_docs.extend(random_docs)
                    
                    random.shuffle(suggested_docs)
                    suggested_docs = suggested_docs[:3]
                else:
                    suggested_docs_ids = db.session.query(Document.id).filter(or_(Document.last_viewed_date.is_(None), Document.engagement_level.is_(None), Document.context_event.is_(None))).limit(20).all()
                    if suggested_docs_ids:
                        random_id_tuples = random.sample(suggested_docs_ids, min(len(suggested_docs_ids), 3))
                        random_ids = [r[0] for r in random_id_tuples]
                        suggested_docs = Document.query.filter(Document.id.in_(random_ids)).all()
            else:
                suggested_docs_ids = db.session.query(Document.id).filter(or_(Document.last_viewed_date.is_(None), Document.engagement_level.is_(None), Document.context_event.is_(None))).limit(20).all()
                if suggested_docs_ids:
                    random_id_tuples = random.sample(suggested_docs_ids, min(len(suggested_docs_ids), 3))
                    random_ids = [r[0] for r in random_id_tuples]
                    suggested_docs = Document.query.filter(Document.id.in_(random_ids)).all()

        except Exception as suggest_err:
            print(f"WARNING: Cannot get suggested docs based on goal: {suggest_err}")
            traceback.print_exc()
            suggested_docs_ids = db.session.query(Document.id).filter(or_(Document.last_viewed_date.is_(None), Document.engagement_level.is_(None), Document.context_event.is_(None))).limit(20).all()
            if suggested_docs_ids:
                random_id_tuples = random.sample(suggested_docs_ids, min(len(suggested_docs_ids), 3))
                random_ids = [r[0] for r in random_id_tuples]
                suggested_docs = Document.query.filter(Document.id.in_(random_ids)).all()
        
        for doc in documents_on_page:
            doc_dict = {
                'id': doc.id,
                'filename': doc.filename,
                'filename_normalized': doc.filename_normalized,
                'doc_type': doc.doc_type,
                'category': doc.category,
                'keywords': doc.keywords,
                'last_viewed_date': doc.last_viewed_date.isoformat() if doc.last_viewed_date else None,
                'engagement_level': doc.engagement_level,
                'context_event': doc.context_event,
                'is_goal_related': getattr(doc, 'is_goal_related', False)
            }
            documents_data_for_js.append(doc_dict)

    except Exception as db_err:
        flash(f"Lỗi khi truy vấn dữ liệu: {db_err}", "danger")
        print(f"ERROR querying database: {db_err}")
        traceback.print_exc()
        
    try:
        version = int(time.time())
        return render_template('index.html', pagination=pagination_data, documents=documents_on_page, documents_data_for_js=documents_data_for_js, categories=available_categories, search_query=search_query, category_filter=category_filter, default_category=fp.DEFAULT_CATEGORY, suggested_docs=suggested_docs,today=date.today(), version=version, timedelta=timedelta)
    except Exception as render_err:
        flash(f"Lỗi nghiêm trọng khi hiển thị trang: {render_err}", "danger")
        print(f"ERROR rendering template: {render_err}")
        traceback.print_exc()
        return "Đã xảy ra lỗi nghiêm trọng khi tải trang.", 500
    
@app.route('/upload', methods=['GET', 'POST'], endpoint='upload_file')
def upload_file_route():
    if request.method == 'POST':
        current_user = User.query.first()
        upload_type = request.form.get('upload_type')
        doc_to_save = None
        learning_goal_form = request.form.get('learning_goal', '').strip()
        deadline_str = request.form.get('deadline', '').strip()
        deadline_date = None
        if deadline_str:
            try:
                deadline_date = datetime.strptime(deadline_str, '%Y-%m-%d').date()
            except ValueError:
                flash('Định dạng ngày không hợp lệ.', 'warning')

        try:
            # Xử lý upload cho từng loại file
            if upload_type == 'file':
                file = request.files.get('document_file')
                if not file or file.filename == '':
                    raise ValueError('Chưa chọn file nào.')
                if not allowed_file(file.filename):
                    raise ValueError('Định dạng file không hợp lệ (cần PDF, DOCX, TXT).')
                
                filename = secure_filename(file.filename)
                filepath = os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                if Document.query.filter_by(filepath=filepath).first():
                    raise ValueError(f'File "{filename}" đã tồn tại.')
                file.save(filepath)
                doc_to_save = Document(filename=filename, filepath=filepath, doc_type=filename.rsplit('.', 1)[1].lower())

            elif upload_type == 'image':
                file = request.files.get('document_image')
                if not file or file.filename == '':
                    raise ValueError('Chưa chọn file ảnh nào.')
                if not allowed_image(file.filename):
                    raise ValueError('Định dạng ảnh không hợp lệ.')

                filename = secure_filename(file.filename)
                filepath = os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                if Document.query.filter_by(filepath=filepath).first():
                    raise ValueError(f'Ảnh "{filename}" đã tồn tại.')
                
                file.save(filepath)
                doc_to_save = Document(filename=filename, filepath=filepath, doc_type='image', category="Hình ảnh")

            elif upload_type == 'video':
                file = request.files.get('document_video')
                if not file or file.filename == '':
                    raise ValueError('Chưa chọn file video nào.')
                if not allowed_video(file.filename):
                    raise ValueError('Định dạng video không hợp lệ.')

                filename = secure_filename(file.filename)
                filepath = os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
                if Document.query.filter_by(filepath=filepath).first():
                    raise ValueError(f'Video "{filename}" đã tồn tại.')

                file.save(filepath)
                doc_to_save = Document(filename=filename, filepath=filepath, doc_type='video', category="Video")

            elif upload_type in ['link', 'googledrive_link']:
                url_key = 'document_url' if upload_type == 'link' else 'document_gdrive_link'
                doc_url = request.form.get(url_key)
                if not doc_url:
                    raise ValueError('Chưa nhập URL.')
                if is_google_drive_link(doc_url) and upload_type == 'link':
                    raise ValueError('Vui lòng dùng mục "Link Google Drive" cho link này.')
                if Document.query.filter_by(filepath=doc_url).first():
                    raise ValueError(f'Link "{doc_url}" đã tồn tại.')
                
                category = "Google Drive" if upload_type == 'googledrive_link' else "Link Web"
                doc_to_save = Document(filename=doc_url, filepath=doc_url, category=category, doc_type=upload_type)
            
            else:
                raise ValueError('Loại upload không hợp lệ.')

            # Nếu đã có đối tượng doc_to_save từ một trong các nhánh trên
            if doc_to_save:
                is_relevant = check_document_relevance(doc_to_save, current_user)

                # --- GÁN CÁC THÔNG TIN VÀO ĐỐI TƯỢNG DOCUMENT ---
                doc_to_save.is_goal_related = is_relevant
                if not doc_to_save.category: # Chỉ gán category nếu nó chưa được set (cho trường hợp link)
                    doc_to_save.category = fp.categorize_document(doc_to_save.filename, current_user.ultimate_goal if current_user else None, current_user.role_model_character if current_user else None)
                
                doc_to_save.filename_normalized = normalize_vietnamese(doc_to_save.filename)
                doc_to_save.learning_goal = learning_goal_form if learning_goal_form else None
                doc_to_save.deadline = deadline_date
                
                # --- LƯU VÀO DATABASE ---
                db.session.add(doc_to_save)
                db.session.commit()

                flash(f'Đã tải lên thành công "{doc_to_save.filename}".', 'info')
                return redirect(url_for('view_document', document_id=doc_to_save.id, review='true'))

        except ValueError as ve:
            flash(str(ve), 'warning')
            return redirect(request.url)
        except Exception as e:
            flash(f'Lỗi không xác định khi upload: {e}', 'danger')
            traceback.print_exc()
            return redirect(request.url)

    return render_template('upload.html')

@app.route('/finalize_upload', methods=['POST'])
def finalize_upload():
    action = request.form.get('action')
    temp_filepath = request.form.get('temp_filepath')
    filename = request.form.get('filename')

    if action == 'reject':
        try:
            if os.path.exists(temp_filepath):
                os.remove(temp_filepath)
            flash(f'Đã hủy và xóa tài liệu tạm thời "{filename}".', 'info')
        except Exception as e:
            flash(f'Lỗi khi xóa file tạm: {e}', 'warning')
        return redirect(url_for('index'))

    # Nếu hành động là lưu (vào Focus hoặc Sandbox)
    is_goal_related = (action == 'save_to_focus')
    permanent_filepath = os.path.abspath(os.path.join(app.config['UPLOAD_FOLDER'], filename))
    if os.path.exists(permanent_filepath):
        flash(f'File "{filename}" đã tồn tại. Vui lòng thử lại với tên khác.', 'danger')
        if os.path.exists(temp_filepath):
            os.remove(temp_filepath)
        return redirect(url_for('upload_file'))
        
    try:
        os.rename(temp_filepath, permanent_filepath)
        category = fp.categorize_document(filename)
        doc_to_save = Document(
            filename=filename, 
            filepath=permanent_filepath,
            doc_type=filename.rsplit('.', 1)[1].lower(), 
            category=category,
            is_goal_related=is_goal_related,
            filename_normalized=normalize_vietnamese(filename)
        )
        db.session.add(doc_to_save)
        db.session.commit()

        if is_goal_related:
            flash(f'Đã lưu "{filename}" vào Focus Workspace!', 'success')
        else:
            flash(f'Đã lưu "{filename}" vào Sandbox.', 'success')
            
        return redirect(url_for('view_document', document_id=doc_to_save.id))

    except Exception as e:
        flash(f'Lỗi nghiêm trọng khi lưu file: {e}', 'danger')
        traceback.print_exc()
        if os.path.exists(temp_filepath): 
             os.remove(temp_filepath)
        return redirect(url_for('upload_file'))

@app.route('/timeline', endpoint='study_timeline')
def study_timeline():
    docs_viewed_today_count = 0; docs_summarized_count = 0
    try:
        today_start_utc = datetime.now(timezone.utc).replace(hour=0, minute=0, second=0, microsecond=0)
        today_end_utc = today_start_utc + timedelta(days=1)
        docs_viewed_today_count = db.session.query(Document.id).filter(Document.last_viewed_date.between(today_start_utc, today_end_utc)).count()
        docs_summarized_count = db.session.query(Document.id).filter(Document.user_summary.isnot(None) & (Document.user_summary != '')).count()
    except Exception as e: print(f"Error fetching timeline stats: {e}")

    streak_days = random.randint(0, 20); avg_session_time = random.randint(10, 45); review_queue = []
    try:
        review_docs = get_random_docs(Document, 5)
        for doc in review_docs:
            due_in_days = random.choice([1, 2, 3, 5, 7]); due_date = date.today() + timedelta(days=due_in_days)
            review_queue.append({"id": doc.id, "name": doc.filename, "due": f"sau {due_in_days} ngày"})
    except Exception as e: print(f"Error fetching review queue: {e}")

    achievements = [{"icon": "bi-emoji-smile", "text": "Bắt đầu hành trình!"}]
    if streak_days >= 5: achievements.append({"icon": "bi-fire", "text": f"{streak_days}-Day Streak!"})
    if docs_summarized_count >= 1: achievements.append({"icon": "bi-lightbulb-fill", "text": "Viết tóm tắt đầu tiên!"})

    return render_template('timeline.html', docs_viewed_today=docs_viewed_today_count, total_docs_summarized=docs_summarized_count, current_streak=streak_days, avg_session_minutes=avg_session_time, review_items=review_queue, earned_achievements=achievements, documents_data_for_js=[]) 

@app.route('/download/<int:document_id>')
def download_file(document_id):
    doc = db.session.get(Document, document_id)
    if not doc or doc.doc_type in ['link', 'googledrive_link']: abort(404)
    try: return send_from_directory(app.config['UPLOAD_FOLDER'], doc.filename, as_attachment=True)
    except Exception as e: flash(f"Lỗi tải file: {e}", "danger"); return redirect(url_for('view_document', document_id=document_id))

@app.route('/document/<int:document_id>', methods=['GET', 'POST'])
def view_document(document_id):
    doc = db.session.get(Document, document_id)
    is_new_upload_for_review = request.args.get('review') == 'true'
    if not doc: abort(404)
    if request.method == 'POST':
        form_marker = request.form.get('form_marker')
        try:
            if form_marker == 'update_context':
                doc.engagement_level = request.form.get('engagement_level')
                doc.context_event = request.form.get('context_event')
                doc.custom_note = request.form.get('custom_note')
                
                doc.learning_goal = request.form.get('learning_goal', '').strip() or None
                deadline_str = request.form.get('deadline', '').strip()
                if deadline_str:
                    try:
                        doc.deadline = datetime.strptime(deadline_str, '%Y-%m-%d').date()
                    except ValueError:
                        flash('Định dạng ngày cho thời hạn không hợp lệ.', 'warning')
                else:
                    doc.deadline = None   
                db.session.commit()
                flash('Đã cập nhật thông tin ngữ cảnh.', 'success')
            elif form_marker == 'submit_summary':
                user_summary = request.form.get('user_summary', '').strip()
                if user_summary:
                    doc.user_summary = user_summary
                    doc.ai_topic_label = fp.categorize_document([user_summary])
                    db.session.commit()
                    flash(f'Đã lưu tóm tắt của bạn. AI gợi ý: {doc.ai_topic_label}', 'info')
                else: flash('Bạn chưa nhập tóm tắt.', 'warning')
            else: flash('Yêu cầu không hợp lệ.', 'warning')
        except Exception as e:
            db.session.rollback()
            flash(f'Lỗi khi xử lý yêu cầu: {e}', 'danger')
        return redirect(url_for('view_document', document_id=document_id))
            
    physical_filepath_to_check = None
    if doc.doc_type not in ['link', 'googledrive_link']:
        try:
            upload_dir = os.path.abspath(app.config['UPLOAD_FOLDER'])
            potential_paths = [os.path.join(upload_dir, secure_filename(doc.filename)), os.path.join(upload_dir, doc.filename)]
            for path in potential_paths:
                if os.path.exists(path): physical_filepath_to_check = path; break
        except Exception as e: print(f"ERROR constructing file path: {e}")
            
    try:
        doc.last_viewed_date = datetime.now(timezone.utc)
        if doc.doc_type in ['pdf', 'docx', 'txt', 'file'] and not doc.extracted_content and physical_filepath_to_check:
            temp_content = fp.extract_text(physical_filepath_to_check)
            doc.extracted_content = temp_content if temp_content and "[Lỗi" not in temp_content else "[File trống hoặc không có nội dung văn bản]"
        db.session.commit()
    except Exception as e: db.session.rollback(); print(f"ERROR updating doc on view: {e}")

    extracted_text_content_for_view = None
    if doc.extracted_content and "[Lỗi" not in doc.extracted_content:
        if '\n' not in doc.extracted_content:
            temp_text = doc.extracted_content
            temp_text = re.sub(r':\s*', ':\n\n', temp_text)
            temp_text = re.sub(r'([.?!])\s+', r'\1\n', temp_text)
            extracted_text_content_for_view = temp_text
        else:
            extracted_text_content_for_view = doc.extracted_content
    
    related_docs = []
    if doc.category and doc.category != fp.DEFAULT_CATEGORY:
        try:
            related_docs = get_random_docs(Document.query.filter(and_(Document.category == doc.category, Document.id != doc.id)), 5)
        except Exception as e: print(f"Error finding related docs: {e}")

    available_categories = sorted(list(fp.CATEGORY_KEYWORDS.keys()))
    if fp.DEFAULT_CATEGORY not in available_categories:
        available_categories.append(fp.DEFAULT_CATEGORY)
        available_categories.sort()

    return render_template('view_document.html', doc=doc, is_new_upload_for_review=is_new_upload_for_review, extracted_content=extracted_text_content_for_view, related_docs=related_docs,  timedelta=timedelta, categories=available_categories, default_category=fp.DEFAULT_CATEGORY)

@app.route('/delete/<int:document_id>', methods=['POST'])
def delete_document(document_id):
    doc = db.session.get(Document, document_id)
    if not doc:
        flash("Tài liệu không tồn tại.", "warning")
        return redirect(url_for('index'))
    filepath_to_delete = None
    if doc.doc_type not in ['link', 'googledrive_link']:
        filepath_to_delete = os.path.join(app.config['UPLOAD_FOLDER'], doc.filename)
    filename_for_flash = doc.filename
    try:
        db.session.delete(doc)
        db.session.commit()
        if filepath_to_delete and os.path.exists(filepath_to_delete):
            os.remove(filepath_to_delete)
        flash(f'Đã xóa thành công tài liệu "{filename_for_flash}".', 'success')

    except Exception as e:
        db.session.rollback()
        flash(f'Lỗi khi xóa tài liệu: {e}', 'danger')
        print(f"ERROR deleting document {document_id}: {e}")
        traceback.print_exc()

    return redirect(url_for('index'))

@app.route('/edit_category/<int:document_id>', methods=['POST'])
def edit_category(document_id):
    doc = db.session.get(Document, document_id)
    new_cat = request.form.get('new_category')
    valid_cats = list(fp.CATEGORY_KEYWORDS.keys()) + [fp.DEFAULT_CATEGORY]
    if doc and new_cat in valid_cats:
        try: doc.category = new_cat; db.session.commit(); flash(f'Đã cập nhật danh mục cho "{doc.filename}".', 'success')
        except Exception as e: db.session.rollback(); flash(f'Lỗi cập nhật danh mục: {e}', 'danger')
    else: flash("Tài liệu hoặc danh mục không hợp lệ.", "warning")
    return redirect(request.referrer or url_for('index'))

@app.route('/document/<int:document_id>/toggle_goal_related', methods=['POST'])
def toggle_goal_related(document_id):
    doc = db.session.get(Document, document_id)
    if not doc:
        return jsonify({"error": "Tài liệu không tồn tại."}), 404
    try:
        doc.is_goal_related = not getattr(doc, 'is_goal_related', False)
        db.session.commit()
        return jsonify({
            "message": "Cập nhật thành công!",
            "is_goal_related": doc.is_goal_related,
            "doc_id": doc.id
        })
    except Exception as e:
        db.session.rollback()
        print(f"ERROR toggling goal related status for doc {document_id}: {e}")
        traceback.print_exc()
        return jsonify({"error": "Lỗi server khi cập nhật trạng thái."}), 500
    
@app.route('/network/<int:document_id>')
def network_view(document_id):
    doc = db.session.get(Document, document_id)
    if not doc: abort(404)
    all_docs = Document.query.order_by(Document.uploaded_date.desc()).all()
    return render_template('network.html', doc=doc, all_docs=all_docs)

@app.route('/clarity_report/<int:doc_id>')
def clarity_report(doc_id):
    doc = db.session.get(Document, doc_id)
    if not doc:
        abort(404)

    # --- TẠO DỮ LIỆU GIẢ LẬP CHO DEMO ---
    transcript = []
    workspace_items = WorkspaceItem.query.filter_by(document_id=doc_id).order_by(WorkspaceItem.order).all()
    workspace_tree = build_tree(workspace_items)
    q_nodes = [item for item in workspace_items if not item.children][:3] 

    if 'dientruong' in doc.filename.lower():
        transcript = [
            {"speaker": "ai", "text": "Chào bạn, hãy bắt đầu ôn tập về điện trường. Câu 1: Điện trường là gì?", "status": "neutral", "node_id": None},
            {"speaker": "user", "text": "Điện trường là một dạng vật chất tồn tại xung quanh điện tích và tác dụng lực điện lên các điện tích khác đặt trong nó.", "status": "correct", "node_id": q_nodes[0].id if q_nodes else None},
            {"speaker": "ai", "text": "Chính xác! Câu 2: Vector cường độ điện trường E có đặc điểm gì?", "status": "neutral", "node_id": None},
            {"speaker": "user", "text": "Vector E có phương, chiều và độ lớn.", "status": "incorrect", "node_id": q_nodes[1].id if len(q_nodes) > 1 else None},
            {"speaker": "ai", "text": "Câu trả lời của bạn chưa đầy đủ. Cần nêu rõ hơn về phương chiều trùng với lực điện, và công thức độ lớn. Câu 3: Nguyên lý chồng chất điện trường phát biểu thế nào?", "status": "neutral", "node_id": None},
            {"speaker": "user", "text": "Là tổng hợp các vector điện trường thành phần.", "status": "correct", "node_id": q_nodes[2].id if len(q_nodes) > 2 else None},
        ]
    else: 
        transcript = [
            {"speaker": "ai", "text": "Chào bạn, sẵn sàng ôn tập chưa? Hãy cho tôi biết khái niệm chính của tài liệu này.", "status": "neutral", "node_id": None},
            {"speaker": "user", "text": "Khái niệm chính là về việc áp dụng mô hình kinh doanh mới.", "status": "correct", "node_id": q_nodes[0].id if q_nodes else None},
            {"speaker": "ai", "text": "Tốt lắm! Yếu tố quan trọng thứ hai là gì?", "status": "neutral", "node_id": None},
            {"speaker": "user", "text": "Tôi không nhớ rõ lắm.", "status": "incorrect", "node_id": q_nodes[1].id if len(q_nodes) > 1 else None},
        ]

    return render_template('clarity_report.html', doc=doc, transcript=transcript, workspace_tree=workspace_tree)

# --- User & Profile Routes ---
@app.route('/profile_setup', methods=['POST'])
def handle_profile_setup_post():
    # SỬA LỖI: Đảm bảo chúng ta định nghĩa và sử dụng biến 'user' một cách nhất quán
    user = User.query.first()

    if not user:
        try:
            # Nếu không có user, tạo một user demo
            new_user = User(username='demo_user', email='demo@example.com',
                            password_hash=generate_password_hash('password'))
            db.session.add(new_user)
            db.session.commit()
            user = new_user # Gán user vừa tạo cho biến 'user'
        except Exception as e:
            flash(f"Lỗi khi tạo người dùng demo: {e}", "danger")
            return redirect(url_for('index'))

    # Lấy dữ liệu từ form
    ultimate_goal = request.form.get('ultimate_goal')
    role_model_character = request.form.get('role_model_character')
    selected_avatar = request.form.get('selected_avatar')
    workspace_color_theme = request.form.get('workspace_color_theme')
    
    # Cập nhật các trường thông tin cho đối tượng 'user'
    user.specific_study_goal = request.form.get('specific_study_goal')
    user.expected_completion_time = request.form.get('expected_completion_time')
    preferred_content_types = request.form.getlist('preferred_content_types')
    user.preferred_content_types = json.dumps(preferred_content_types) if preferred_content_types else None
    personal_learning_challenges = request.form.getlist('personal_learning_challenges')
    user.personal_learning_challenges = json.dumps(personal_learning_challenges) if personal_learning_challenges else None
    studyvault_expectations = request.form.getlist('studyvault_expectations')
    user.studyvault_expectations = json.dumps(studyvault_expectations) if studyvault_expectations else None
    user.ultimate_goal = ultimate_goal
    user.role_model_character = role_model_character
    user.selected_avatar = selected_avatar
    user.workspace_color_theme = workspace_color_theme

    try:
        db.session.commit()
        flash('Thiết lập hồ sơ thành công!', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Lỗi khi lưu hồ sơ: {e}', 'danger')
        print(f"ERROR saving profile: {e}")
        traceback.print_exc() 

    return redirect(url_for('index'))

@app.route('/activate_short_term_mode', methods=['POST'])
def activate_short_term_mode():
    current_user = User.query.first()
    if not current_user:
        flash("Bạn cần đăng nhập để kích hoạt chế độ này.", "danger")
        return jsonify({"success": False, "error": "Bạn cần đăng nhập để kích hoạt chế độ này."}), 401

    study_duration = request.form.get('study_duration', type=int)
    study_focus_keywords = request.form.get('study_focus_keywords')
    study_intensity = request.form.get('study_intensity')

    if not study_duration or study_duration < 1 or study_duration > 7:
        return jsonify({"success": False, "error": "Thời gian học không hợp lệ. Vui lòng chọn từ 1 đến 7 ngày."}), 400

    try:
        current_user.short_term_mode_active = True
        current_user.short_term_mode_end_date = date.today() + timedelta(days=study_duration)
        current_user.short_term_mode_focus_keywords = study_focus_keywords
        current_user.short_term_mode_intensity = study_intensity
        current_user.short_term_mode_last_used = date.today() 

        db.session.commit()
        
        return jsonify({
            "success": True,
            "message": f"SHORT-TERM MODE đã được kích hoạt trong {study_duration} ngày! Hãy tập trung học tập!",
        }), 200
    except Exception as e:
        db.session.rollback()
        print(f"ERROR activating short-term mode: {e}")
        traceback.print_exc()
        return jsonify({"success": False, "error": f"Lỗi khi kích hoạt SHORT-TERM MODE: {e}"}), 500
    
@app.route('/deactivate_short_term_mode', methods=['POST'])
def deactivate_short_term_mode():
    current_user = User.query.first()
    if not current_user:
        return jsonify({"error": "Bạn cần đăng nhập để thực hiện hành động này."}), 401

    try:
        current_user.short_term_mode_active = False
        # current_user.short_term_mode_end_date = None # Có thể giữ lại để hiển thị thông tin kết thúc
        # current_user.short_term_mode_focus_keywords = None
        # current_user.short_term_mode_intensity = None
        db.session.commit()
        return jsonify({"message": "SHORT-TERM MODE đã được vô hiệu hóa."}), 200
    except Exception as e:
        db.session.rollback()
        print(f"ERROR deactivating short-term mode: {e}")
        traceback.print_exc()
        return jsonify({"error": f"Lỗi server: {e}"}), 500
    
    # --- API Endpoints ---
@app.route('/document/<int:doc_id>/workspace', methods=['GET'])
def get_workspace_data(doc_id):
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({"error": "Tài liệu không tồn tại."}), 404

    try:
        all_items = WorkspaceItem.query.filter_by(document_id=doc_id).order_by(WorkspaceItem.order).all()
        workspace_tree = build_tree(all_items)
        return jsonify(workspace_tree)

    except Exception as e:
        print(f"ERROR in get_workspace_data for doc {doc_id}: {e}")
        traceback.print_exc()
        return jsonify({"error": "Lỗi server khi lấy dữ liệu workspace."}), 500
    
@app.route('/document/<int:doc_id>/workspace_items', methods=['POST'])
def create_workspace_item(doc_id):
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({"error": "Tài liệu không tồn tại."}), 404
    data = request.get_json()
    title = data.get('title', '').strip()
    if not title:
        return jsonify({"error": "Tiêu đề không được để trống."}), 400
    content = data.get('content', '')
    parent_id = data.get('parent_id')
    try:
        if parent_id:
            max_order = db.session.query(sql_func.max(WorkspaceItem.order)).filter_by(parent_id=parent_id).scalar() or 0
        else:
            max_order = db.session.query(sql_func.max(WorkspaceItem.order)).filter_by(document_id=doc_id, parent_id=None).scalar() or 0

        new_item = WorkspaceItem(
            title=title, content=content, order=max_order + 1,
            document_id=doc_id, parent_id=parent_id if parent_id else None
        )
        db.session.add(new_item)
        db.session.commit()

        return jsonify({ "id": new_item.id, "title": new_item.title, "content": new_item.content, "order": new_item.order, "parent_id": new_item.parent_id, "children": [] }), 201

    except Exception as e:
        db.session.rollback()
        print(f"ERROR creating workspace_item for doc {doc_id}: {e}")
        return jsonify({"error": "Lỗi server khi tạo mục."}), 500

@app.route('/workspace_item/<int:item_id>/user_content', methods=['POST'])
def save_workspace_item_user_content(item_id):
    item = db.session.get(WorkspaceItem, item_id)
    if not item:
        return jsonify({"error": "Mục không tồn tại."}), 404
    data = request.get_json()
    user_content = data.get('user_content', '')
    try:
        item.user_content = user_content
        db.session.commit()
        questions = []
        user_content_lower = user_content.lower()
        if 'điện trường' in user_content_lower:
            questions.extend([
                {"id": 1, "q": "Điện trường là gì và tính chất cơ bản nhất của nó là gì?", "type": "Định nghĩa"},
                {"id": 2, "q": "Vector cường độ điện trường tại một điểm được định nghĩa như thế nào? Nêu đặc điểm (điểm đặt, phương, chiều, độ lớn).", "type": "Khái niệm"},
                {"id": 3, "q": "Phân biệt điện trường do điện tích điểm Q dương và điện tích điểm Q âm gây ra.", "type": "So sánh"}
            ])
        elif 'công thức' in user_content_lower and 'e =' in user_content_lower:
            questions.append({"id": 4, "q": "Áp dụng công thức E = k|q|/r² để tính E tại một điểm M cách q=5nC một khoảng 10cm.", "type": "Vận dụng"})
        
       
        if not questions:
            questions.append({"id": 0, "q": "Nội dung chưa đủ để tạo câu hỏi, hãy thử phân tích sâu hơn về một khái niệm cụ thể.", "type": "Gợi ý"})

        return jsonify({"message": "Lưu thành công!", "questions": questions})

    except Exception as e:
        db.session.rollback()
        return jsonify({"error": "Lỗi server khi lưu nội dung."}), 500
    
@app.route('/workspace_item/<int:item_id>/ai_suggestion', methods=['POST'])
def get_ai_suggestion(item_id):
    if not db.session.get(WorkspaceItem, item_id):
        return jsonify({"error": "Mục không tồn tại."}), 404
        
    data = request.get_json()
    user_content = data.get('user_content', '').lower()

    suggestion = ""
    if 'điện trường' in user_content:
        suggestion = (
            "Thử giải thích bằng một phép so sánh đơn giản xem sao nhé! "
            "Ví dụ: 'Hãy tưởng tượng bạn đứng gần một cái lò sưởi, dù không chạm vào nhưng vẫn thấy ấm. Cái ấm đó lan ra từ lò - giống như điện trường lan ra từ một vật mang điện. "
            "Khi bạn đặt một vật mang điện (như cực pin) ở đâu đó, xung quanh nó sẽ xuất hiện một vùng vô hình. Nếu đặt một vật khác cũng có điện vào vùng đó, nó sẽ bị hút hoặc đẩy. Vùng vô hình ấy chính là điện trường.'"
        )
    elif 'tĩnh điện' in user_content or 'lực điện' in user_content:
        suggestion = (
            "Bạn có thể dùng câu chuyện về chiếc bóng bay để giải thích! "
            "Ví dụ: 'Khi bạn chà xát một quả bóng bay vào tóc, những hạt bụi nhỏ li ti từ tóc sẽ dính sang bóng và ngược lại. "
            "Những hạt bụi này có khả năng hút hoặc đẩy nhau. Nếu chúng &quot;thích&quot; nhau, chúng sẽ kéo lại gần; nếu &quot;ghét&quot; nhau, chúng sẽ đẩy nhau ra. "
            "Kết luận: Hạt bụi có khả năng hút đẩy nhau gọi là <b>điện tích</b>. &quot;Tình cảm&quot; giữa các hạt bụi (hút hoặc đẩy) chính là <b>lực điện</b>.'"
        )
    else:
        generic_suggestions = [
            "Bạn có thể diễn giải lại một cách đơn giản hơn được không?",
            "Thử mở rộng thêm lý thuyết liên quan xem sao.",
            "Hãy thử liên kết khái niệm này với một ví dụ trong thực tế."
        ]
        suggestion = random.choice(generic_suggestions)

    return jsonify({"suggestion": suggestion})

@app.route('/workspace_item/<int:item_id>/labels', methods=['PUT'])
def update_workspace_item_labels(item_id):
    item = db.session.get(WorkspaceItem, item_id)
    if not item:
        return jsonify({"error": "Mục không tồn tại."}), 404

    data = request.get_json()
    if not data:
        return jsonify({"error": "Dữ liệu không hợp lệ."}), 400

    try:
        item.importance = data.get('importance', item.importance)
        item.learning_role = data.get('learning_role', item.learning_role)
        item.difficulty = data.get('difficulty', item.difficulty)
        db.session.commit()

        return jsonify({
            "id": item.id,
            "title": item.title,
            "importance": item.importance,
            "learning_role": item.learning_role,
            "difficulty": item.difficulty
        })

    except Exception as e:
        db.session.rollback()
        print(f"ERROR updating labels for item {item_id}: {e}")
        return jsonify({"error": "Lỗi server khi cập nhật nhãn."}), 500

@app.route('/tokenize_content', methods=['POST'])
def tokenize_content():
    try:
        data = request.get_json()
        text_content = data.get('text', '')
        category = data.get('category', 'Chưa phân loại')

        if not text_content:
            return jsonify({"error": "Không có nội dung để phân tích."}), 400

        normalized_content = normalize_vietnamese(text_content)
        found_keywords = set()

        keywords_for_category = fp.CATEGORY_KEYWORDS.get(category, [])
        for keyword in keywords_for_category:
            if re.search(r'\b' + re.escape(normalize_vietnamese(keyword)) + r'\b', normalized_content):
                found_keywords.add(keyword.capitalize())

        
        if not found_keywords:
            return jsonify({"keywords": ["Không tìm thấy từ khóa chuyên ngành nào."]})

        return jsonify({"keywords": sorted(list(found_keywords))})

    except Exception as e:
        print(f"ERROR in /tokenize_content: {e}")
        traceback.print_exc()
        return jsonify({"error": "Lỗi server khi phân tích từ khóa."}), 500
    
@app.route('/get_recall_data', methods=['GET'])
def get_recall_data():
    TOTAL_QUESTIONS_TO_RETURN = 3
    all_recall_items = []
    doc_ids_used = set()
    try:
        docs_with_summary = Document.query.filter(Document.user_summary.isnot(None), Document.user_summary != '').all()
        if docs_with_summary:
            random.shuffle(docs_with_summary)
            for doc in docs_with_summary:
                if len(all_recall_items) >= 2: break
                question_text = f"Hãy nêu lại ý chính/định nghĩa bạn đã tóm tắt cho tài liệu '{doc.filename}'."
                answer_text = doc.user_summary
                all_recall_items.append({ "q": question_text, "a": answer_text, "cat": doc.category or "Từ tóm tắt", "source_doc_id": doc.id, "type": "definition_recall" })
                doc_ids_used.add(doc.id)
    except Exception as e:
        print(f"WARNING: Lỗi khi tạo câu hỏi từ tóm tắt. Lỗi: {e}")

    needed = TOTAL_QUESTIONS_TO_RETURN - len(all_recall_items)
    if needed > 0:
        try:
            docs_with_content = Document.query.filter(Document.id.notin_(list(doc_ids_used)), Document.extracted_content.isnot(None), Document.extracted_content != '', Document.extracted_content.notlike('[%')).all()
            if docs_with_content:
                random.shuffle(docs_with_content)
                for doc in docs_with_content:
                    if len(all_recall_items) >= TOTAL_QUESTIONS_TO_RETURN: break
                    try:
                        for _ in range(random.randint(1, 2)):
                            q = create_fill_in_the_blank_question(doc.extracted_content)
                            if q and len(all_recall_items) < TOTAL_QUESTIONS_TO_RETURN:
                                all_recall_items.append({ "q": q["q"], "a": q["a"], "cat": doc.category or "Từ tài liệu", "source_doc_id": doc.id, "type": "fill_blank" })
                                doc_ids_used.add(doc.id)
                            else:
                                break
                    except Exception as q_err:
                        print(f"WARNING: Không thể tạo câu hỏi điền vào chỗ trống cho doc_id {doc.id}. Lỗi: {q_err}")
        except Exception as e:
            print(f"WARNING: Lỗi khi truy vấn tài liệu có nội dung. Lỗi: {e}")

    needed = TOTAL_QUESTIONS_TO_RETURN - len(all_recall_items)
    if needed > 0:
        defaults = [
            {"q": "Trong Python, `list` và `tuple` khác nhau thế nào?", "a": "List có thể thay đổi (mutable), tuple thì không (immutable).", "cat": "Lập trình", "type": "default"},
            {"q": "Lệnh `git push` dùng để làm gì?", "a": "Đẩy các commit từ local repository lên remote repository.", "cat": "Lập trình", "type": "default"},
            {"q": "API là viết tắt của cụm từ gì?", "a": "Application Programming Interface (Giao diện lập trình ứng dụng)", "cat": "Lập trình", "type": "default"},
            
            {"q": "GDP là viết tắt của gì?", "a": "Gross Domestic Product (Tổng sản phẩm quốc nội)", "cat": "Kinh tế học", "type": "default"},
            {"q": "Lạm phát là gì?", "a": "Sự tăng mức giá chung một cách liên tục của hàng hóa và dịch vụ theo thời gian, và sự mất giá trị của một loại tiền tệ nào đó.", "cat": "Kinh tế học", "type": "default"},

            {"q": "Đạo hàm của hàm số f(x) = x² là gì?", "a": "f'(x) = 2x", "cat": "Toán học", "type": "default"},
            {"q": "Số Pi (π) có giá trị xấp xỉ bằng bao nhiêu?", "a": "3.14159", "cat": "Toán học", "type": "default"},

            {"q": "Định luật Vạn vật hấp dẫn do ai phát minh?", "a": "Isaac Newton", "cat": "Vật Lý", "type": "default"},
            
            {"q": "Công thức hóa học của nước là gì?", "a": "H₂O", "cat": "Hóa học", "type": "default"},

            {"q": "Sự khác biệt giữa 'your' và 'you're' trong tiếng Anh là gì?", "a": "'Your' là tính từ sở hữu, 'you're' là viết tắt của 'you are'.", "cat": "Ngoại ngữ", "type": "default"},

            {"q": "Năm chữ cái trong phương pháp đặt mục tiêu S.M.A.R.T. là viết tắt của những từ gì?", "a": "Specific, Measurable, Achievable, Relevant, Time-bound (Cụ thể, Đo lường được, Khả thi, Liên quan, Có thời hạn)", "cat": "Kỹ năng mềm", "type": "default"},
            
            {"q": "Học có giám sát (Supervised Learning) khác gì với Học không giám sát (Unsupervised Learning)?", "a": "Học có giám sát dùng dữ liệu đã được gán nhãn, trong khi học không giám sát dùng dữ liệu không gán nhãn.", "cat": "AI/ML", "type": "default"}
        ]
        all_recall_items.extend(get_unique_random_elements(defaults, needed))
    random.shuffle(all_recall_items)
    final_questions = all_recall_items[:TOTAL_QUESTIONS_TO_RETURN]

    if not final_questions: 
        return jsonify([{"q": "Hiện không có câu hỏi nào. Hãy tải và tóm tắt thêm tài liệu nhé!", "a": "", "cat": "Hệ thống", "type": "error"}])
    
    return jsonify(final_questions)

@app.route('/ai_chat_converse', methods=['POST'])
def ai_chat_converse():
    try:
        data = request.get_json()
        user_message = data.get('message', '').strip()
        if not user_message:
            return jsonify({"error": "Không có tin nhắn nào được gửi."}), 400

        ai_reply = ""
        user_message_normalized = normalize_vietnamese(user_message)
        
        if any(word in user_message_normalized for word in ["tam biet", "bye", "ket thuc", "cam on"]):
            return jsonify({"reply": random.choice(["Rất vui được hỗ trợ bạn! Hẹn gặp lại nhé!", "Chúc bạn học tốt! Nếu cần gì cứ gọi tôi."])})
        
        is_summary_request = "tom tat" in user_message_normalized
        
        search_query = user_message_normalized.replace("tom tat", "").strip()
        search_query_with_spaces = search_query.replace("_", " ")
        
        found_doc = None
        if len(search_query) > 2:
            found_doc = Document.query.filter(
                sql_func.replace(Document.filename_normalized, '_', ' ').ilike(f"%{search_query_with_spaces}%")
            ).first()
            
        if found_doc:
            if is_summary_request:
                reply = f"Đây là tóm tắt của bạn cho tài liệu '{found_doc.filename}':\n\n\"{found_doc.user_summary}\"" if found_doc.user_summary else f"Rất tiếc, bạn chưa có tóm tắt nào cho tài liệu '{found_doc.filename}'."
            else:
                reply = f"Tôi đã tìm thấy tài liệu '{found_doc.filename}'. Bạn muốn hỏi gì về nó, hay muốn xem tóm tắt?"
        else:
            if is_summary_request:
                reply = "Bạn muốn tóm tắt tài liệu nào? Vui lòng cho tôi biết tên đầy đủ và chính xác hơn nhé."
            elif any(word in user_message_normalized for word in ["chao", "hello", "hi"]):
                random_doc = get_random_docs(Document, 1)
                if random_doc:
                    reply = f"Chào bạn! Bạn có muốn thảo luận về tài liệu '{random_doc[0].filename}' hoặc một chủ đề nào khác không?"
                else:
                    reply = "Chào bạn! Hôm nay bạn muốn thảo luận về chủ đề nào?"
            else:
                reply = f"Tôi không tìm thấy tài liệu nào khớp với '{user_message}'. Bạn có thể thử lại với tên khác không?"
        return jsonify({"reply": reply})
    except Exception as e:
        print(f"ERROR in /ai_chat_converse: {e}")
        traceback.print_exc()
        return jsonify({"error": "Xin lỗi, tôi đang gặp chút sự cố và không thể trả lời ngay. Bạn thử lại sau nhé."}), 500

@app.route('/document/<int:doc_id>/win_criteria', methods=['PUT'])
def update_win_criteria(doc_id):
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({"error": "Tài liệu không tồn tại."}), 404
    data = request.get_json()
    if not data:
        return jsonify({"error": "Dữ liệu không hợp lệ."}), 400

    try:
        doc.win_criteria_description = data.get('description', doc.win_criteria_description)
        target_score_str = data.get('target_score')
        if target_score_str is not None and target_score_str != '':
            doc.target_score = int(target_score_str)
        else:
            doc.target_score = None 
        db.session.commit()
        return jsonify({
            "message": "Cập nhật mục tiêu thành công!",
            "win_criteria_description": doc.win_criteria_description,
            "target_score": doc.target_score
        })

    except (ValueError, TypeError):
        return jsonify({"error": "Điểm mục tiêu phải là một con số."}), 400
    except Exception as e:
        db.session.rollback()
        print(f"ERROR updating win criteria for doc {doc_id}: {e}")
        return jsonify({"error": "Lỗi server khi cập nhật mục tiêu."}), 500

@app.route('/document/<int:doc_id>/result', methods=['POST'])
def log_actual_result(doc_id):
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({"error": "Tài liệu không tồn tại."}), 404

    data = request.get_json()
    if not data or 'actual_score' not in data:
        return jsonify({"error": "Dữ liệu không hợp lệ, thiếu điểm thực tế."}), 400    
    try:
        actual_score_str = data.get('actual_score')
        if actual_score_str is not None and actual_score_str != '':
             doc.actual_score = int(actual_score_str)
        else:
            doc.actual_score = None 
        db.session.commit()
        return jsonify({
            "message": "Ghi nhận kết quả thành công!",
            "actual_score": doc.actual_score
        })

    except (ValueError, TypeError):
        return jsonify({"error": "Điểm thực tế phải là một con số."}), 400
    except Exception as e:
        db.session.rollback()
        print(f"ERROR logging result for doc {doc_id}: {e}")
        return jsonify({"error": "Lỗi server khi ghi nhận kết quả."}), 500

@app.route('/api/generate_questions_from_summary', methods=['POST'])
def generate_questions_from_summary():
    data = request.get_json()
    summary_text = data.get('summary_text', '').lower()
    questions = []
    if 'điện trường' in summary_text and 'tác dụng lực' in summary_text:
        questions = [
            {"id": 1, "q": "Điện trường là gì và tính chất cơ bản nhất của nó là gì?", "type": "Định nghĩa"},
            {"id": 2, "q": "Làm thế nào để có thể phát hiện sự tồn tại của điện trường tại một điểm?", "type": "Khái niệm"},
            {"id": 3, "q": "Nếu một electron (điện tích âm) bay vào một điện trường, nó sẽ di chuyển cùng chiều hay ngược chiều với chiều của vector cường độ điện trường?", "type": "Vận dụng"}
        ]
    elif 'lực điện' in summary_text:
         questions = [
            {"id": 1, "q": "Phát biểu công thức của lực điện (lực Coulomb).", "type": "Công thức"},
            {"id": 2, "q": "Lực điện là lực hút hay đẩy khi hai điện tích trái dấu? Cùng dấu?", "type": "Khái niệm"}
        ]
    else:
        questions.append({"id": 0, "q": "Nội dung tóm tắt chưa đủ rõ ràng để AI có thể tạo câu hỏi. Vui lòng cung cấp thêm chi tiết về các định nghĩa hoặc công thức.", "type": "Thông báo"})

    return jsonify(questions)

@app.route('/document/<int:doc_id>/objectives', methods=['POST'])
def add_objective(doc_id):
    data = request.get_json()
    description = data.get('description', '').strip()
    parent_id = data.get('parent_id')
    if not description:
        return jsonify({"error": "Nội dung không được để trống"}), 400

    new_obj = LearningObjective(description=description, document_id=doc_id, parent_id=parent_id)
    db.session.add(new_obj)
    db.session.commit()
    return jsonify({"id": new_obj.id, "description": new_obj.description, "is_completed": new_obj.is_completed}), 201

@app.route('/objective/<int:obj_id>', methods=['DELETE'])
def delete_objective(obj_id):
    obj = db.session.get(LearningObjective, obj_id)
    if not obj:
        return jsonify({"error": "Mục tiêu không tồn tại"}), 404
    db.session.delete(obj)
    db.session.commit()
    return jsonify({"message": "Xóa thành công"}), 200

@app.route('/objective/<int:obj_id>/toggle', methods=['PUT'])
def toggle_objective(obj_id):
    obj = db.session.get(LearningObjective, obj_id)
    if not obj:
        return jsonify({"error": "Mục tiêu không tồn tại"}), 404
    obj.is_completed = not obj.is_completed
    db.session.commit()
    return jsonify({"is_completed": obj.is_completed}), 200

@app.route('/document/<int:doc_id>/objectives_tree', methods=['GET'])
def get_objectives_tree(doc_id):
    if not db.session.get(Document, doc_id):
        return jsonify({"error": "Tài liệu không tồn tại."}), 404
    all_objectives = LearningObjective.query.filter_by(document_id=doc_id).all()
    tree = FileProcessor.build_objectives_tree(all_objectives) 
    return jsonify(tree)

@app.route('/api/document/<int:document_id>/network')
def get_document_network_mermaid(document_id):
    doc = db.session.get(Document, document_id)
    if not doc:
        return "Tài liệu không tồn tại", 404

    all_items = WorkspaceItem.query.filter_by(document_id=document_id).order_by(WorkspaceItem.order).all()
    mermaid_data = generate_mermaid_graph(all_items)
    return mermaid_data, 200, {'Content-Type': 'text/plain; charset=utf-8'}

@app.route('/api/workspace_item/<int:item_id>')
def get_workspace_item_details(item_id):
    item = db.session.get(WorkspaceItem, item_id)
    if not item:
        return jsonify({"error": "Mục không tồn tại"}), 404

    return jsonify({
        "id": item.id,
        "title": item.title,
        "content": item.content,
        "user_content": item.user_content
    })

@app.route('/api/workspace_item/<int:item_id>/update', methods=['POST'])
def update_workspace_item_content(item_id):
    item = db.session.get(WorkspaceItem, item_id)
    if not item:
        return jsonify({"error": "Mục không tồn tại"}), 404

    data = request.get_json()
    if 'user_content' not in data:
        return jsonify({"error": "Thiếu nội dung"}), 400

    try:
        item.user_content = data['user_content']
        db.session.commit()
        return jsonify({"message": "Lưu thành công!"})
    except Exception as e:
        db.session.rollback()
        return jsonify({"error": f"Lỗi server: {e}"}), 500

@app.route('/api/document/<int:doc_id>/auto_breakdown', methods=['POST'])
def auto_breakdown_document(doc_id):
    doc = db.session.get(Document, doc_id)
    if not doc:
        return jsonify({"error": "Tài liệu không tồn tại"}), 404
    try: 
        WorkspaceItem.query.filter_by(document_id=doc_id).delete()
        db.session.flush()
        doc_filename_lower = doc.filename.lower()

        # === LOGIC ĐÃ SỬA LỖI: Kiểm tra từ khóa cụ thể hơn ===
        
        # Kịch bản 1: Cho file chứa "ai-voice"
        if 'ai-voice' in doc_filename_lower:
            root = WorkspaceItem(title="Công nghệ AI Voice", document_id=doc_id, order=1)
            db.session.add(root)
            db.session.flush()
            stt = WorkspaceItem(title="Speech to Text (STT)", document_id=doc_id, parent_id=root.id, order=1, content="Chuyển đổi giọng nói thành văn bản.")
            tts = WorkspaceItem(title="Text to Speech (TTS)", document_id=doc_id, parent_id=root.id, order=2, content="Chuyển đổi văn bản thành giọng nói.")
            nlu = WorkspaceItem(title="Natural Language Understanding (NLU)", document_id=doc_id, parent_id=root.id, order=3, content="Giúp máy tính hiểu ý nghĩa của ngôn ngữ tự nhiên.")
            db.session.add_all([stt, tts, nlu])
            db.session.flush()
            nlu_child1 = WorkspaceItem(title="Intent Recognition", document_id=doc_id, parent_id=nlu.id, order=1, content="Nhận diện ý định của người dùng (vd: đặt báo thức, hỏi thời tiết).")
            nlu_child2 = WorkspaceItem(title="Entity Extraction", document_id=doc_id, parent_id=nlu.id, order=2, content="Trích xuất các thực thể quan trọng (vd: thời gian, địa điểm).")
            db.session.add_all([nlu_child1, nlu_child2])

        # Kịch bản 2: Cho file chứa "dientruong"
        elif 'dientruong' in doc_filename_lower:
            root = WorkspaceItem(title="Lý thuyết về Điện trường", document_id=doc_id, order=1)
            db.session.add(root)
            db.session.flush()
            child1 = WorkspaceItem(title="Định nghĩa Điện trường", document_id=doc_id, parent_id=root.id, order=1, content="Điện trường là một dạng vật chất tồn tại xung quanh điện tích và tác dụng lực điện lên các điện tích khác đặt trong nó.")
            child2 = WorkspaceItem(title="Vector Cường độ Điện trường (E)", document_id=doc_id, parent_id=root.id, order=2)
            child3 = WorkspaceItem(title="Nguyên lý chồng chất Điện trường", document_id=doc_id, parent_id=root.id, order=3, content="Vector cường độ điện trường tại một điểm bằng tổng các vector cường độ điện trường do từng điện tích điểm gây ra tại điểm đó.")
            db.session.add_all([child1, child2, child3])
            db.session.flush()
            grandchild1 = WorkspaceItem(title="Công thức tính E", document_id=doc_id, parent_id=child2.id, order=1, content="E = k * |Q| / (ε * r²)")
            grandchild2 = WorkspaceItem(title="Đặc điểm của Vector E", document_id=doc_id, parent_id=child2.id, order=2)
            db.session.add_all([grandchild1, grandchild2])

        else:
            root = WorkspaceItem(title=f"Sơ đồ cho '{doc.filename}'", document_id=doc_id, order=1)
            db.session.add(root)
        db.session.commit()
        return jsonify({"message": "Auto Breakdown thành công!"})
    except Exception as e:
        db.session.rollback()
        print(f"Lỗi khi auto breakdown: {e}")
        traceback.print_exc()
        return jsonify({"error": "Lỗi server khi thực hiện breakdown."}), 500
    

@app.route('/api/network/<int:doc_id>/find_and_merge', methods=['POST'])
def find_and_merge_nodes(doc_id):
    if not db.session.get(Document, doc_id):
        return jsonify({"error": "Tài liệu không tồn tại"}), 404

    try:
        nodes = WorkspaceItem.query.filter_by(document_id=doc_id).order_by(WorkspaceItem.id.desc()).all()
        if len(nodes) < 2:
            return jsonify({"error": "Không có đủ node để thực hiện gộp."}), 400
        node_to_be_merged = nodes[0]  
        target_node = nodes[1]       

        return jsonify({
            "message": "AI đã tìm thấy 2 node có thể gộp.",
            "source_node_id": node_to_be_merged.id,
            "target_node_id": target_node.id,
            "source_node_title": node_to_be_merged.title,
            "target_node_title": target_node.title
        })
        
    except Exception as e:
        print(f"Lỗi trong find_and_merge_nodes: {e}")
        traceback.print_exc()
        return jsonify({"error": "Lỗi server khi tìm kiếm node."}), 500
    
@app.route('/api/simplify_check', methods=['POST'])
def check_simplification():
    data = request.get_json()
    text = data.get('text', '').strip()
    if not text:
        return jsonify({"error": "Vui lòng nhập nội dung để kiểm tra."}), 400
    issues = []
    # --- Mô phỏng 5 quy tắc ---
    passive_verbs = re.findall(r'\b(is|was|are|were)\s+(\w+ed)\b', text, re.IGNORECASE)
    if passive_verbs:
        issues.append({"type": "Bị động", "text": " ".join(passive_verbs[0]), "suggestion": "Hãy thử chuyển câu này về thể chủ động để rõ nghĩa hơn."})

    complex_conjunctions = ['because', 'although', 'which', 'while', 'since', 'that']
    if any(f' {conj} ' in text.lower() for conj in complex_conjunctions):
        issues.append({"type": "Câu phức", "suggestion": "Câu của bạn có vẻ chứa nhiều ý. Hãy thử tách thành các câu đơn ngắn gọn hơn."})
    
    verbs = re.findall(r'\b(is|are|was|were|has|have|go|goes|do|does|explain|create)\b', text.lower())
    if len(verbs) > 2:
         issues.append({"type": "Nhiều ý tưởng", "suggestion": "Câu văn có vẻ hơi dài. Hãy đảm bảo mỗi câu chỉ diễn đạt một ý chính."})

    jargon_words = ['vector', 'cường độ', 'chồng chất', 'thuật toán', 'mô hình hóa', 'thực thể']
    found_jargon = [jargon for jargon in jargon_words if jargon in text.lower()]
    if found_jargon:
        issues.append({"type": "Thuật ngữ", "text": ", ".join(found_jargon), "suggestion": "Phát hiện có thuật ngữ chuyên ngành. Bạn có chắc người mới học có thể hiểu được không?"})

    example_words = ['ví dụ', 'vd:', 'chẳng hạn', 'for example', 'e.g.']
    if not any(ex_word in text.lower() for ex_word in example_words):
        issues.append({"type": "Thiếu ví dụ", "suggestion": "Thêm một ví dụ thực tế sẽ giúp người khác dễ hiểu hơn rất nhiều!"})
    
    if not issues:
        words_in_text = [word for word in re.split(r'\s+', text) if len(word) > 4]
        quiz_questions = []
        if len(words_in_text) > 2:
            keyword1 = random.choice(words_in_text)
            remaining_words = [w for w in words_in_text if w != keyword1]
            keyword2 = random.choice(remaining_words) if remaining_words else keyword1
            
            quiz_questions.append({
                "question": f"Trong phần giải thích của bạn, '{keyword1}' có ý nghĩa là gì?",
                "options": ["Định nghĩa đúng ", "Định nghĩa sai 1", "Định nghĩa sai 2"],
                "answer": "Định nghĩa đúng "
            })
            quiz_questions.append({
                "question": f"Mối quan hệ giữa '{keyword1}' và '{keyword2}' là gì?",
                "options": ["Quan hệ đúng ", "Quan hệ sai 1", "Quan hệ sai 2"],
                "answer": "Quan hệ đúng "
            })
            quiz_questions.append({
                "question": "Ý chính của toàn bộ đoạn giải thích là gì?",
                "options": ["Ý chính đúng ", "Ý chính sai 1", "Ý chính sai 2"],
                "answer": "Ý chính đúng "
            })

        return jsonify({
            "original_text": text, 
            "issues": [], 
            "quiz_questions": quiz_questions
        })
    else:
        return jsonify({"original_text": text, "issues": issues})

@app.route('/api/document/<int:document_id>/workspace_for_graph')
def get_workspace_for_graph(document_id):
    doc = db.session.get(Document, document_id)
    if not doc: return jsonify({"error": "Tài liệu không tồn tại"}), 404 
    nodes = WorkspaceItem.query.filter_by(document_id=document_id).all()
    relations = WorkspaceItemRelation.query.filter_by(document_id=document_id).all()
    nodes_data = [{"id": node.id, "title": node.title} for node in nodes]
    relations_data = [{"source_id": rel.source_id, "target_id": rel.target_id, "label": rel.label} for rel in relations]
    for node in nodes:
        if node.parent_id:
            relations_data.append({"source_id": node.parent_id, "target_id": node.id, "label": "là một phần của"})

    return jsonify({"nodes": nodes_data, "relations": relations_data})

@app.route('/api/document/<int:doc_id>/relations', methods=['POST'])
def add_relation(doc_id):
    doc = db.session.get(Document, doc_id)
    if not doc: return jsonify({"error": "Tài liệu không tồn tại"}), 404

    data = request.get_json()
    source_id = data.get('source_id')
    target_id = data.get('target_id')
    label = data.get('label', '')

    source_node = WorkspaceItem.query.filter_by(id=source_id, document_id=doc_id).first()
    target_node = WorkspaceItem.query.filter_by(id=target_id, document_id=doc_id).first()

    if not source_node or not target_node: return jsonify({"error": "Node nguồn hoặc đích không hợp lệ"}), 400

    new_relation = WorkspaceItemRelation(document_id=doc_id, source_id=source_id, target_id=target_id, label=label)
    db.session.add(new_relation)
    db.session.commit()

    return jsonify({"message": "Tạo quan hệ thành công", "relation_id": new_relation.id}), 201

@app.route('/api/network/<int:doc_id>/execute_merge', methods=['POST'])
def execute_merge_nodes(doc_id):
    if not db.session.get(Document, doc_id):
        return jsonify({"error": "Tài liệu không tồn tại"}), 404

    try:
        nodes = WorkspaceItem.query.filter_by(document_id=doc_id).order_by(WorkspaceItem.id.desc()).limit(2).all()
        
        if len(nodes) < 2:
            return jsonify({"error": "Không có đủ node để thực hiện gộp."}), 400
        source_node = nodes[0]
        target_node = nodes[1]
        WorkspaceItem.query.filter_by(parent_id=source_node.id).update({'parent_id': target_node.id})
        WorkspaceItemRelation.query.filter_by(source_id=source_node.id, document_id=doc_id).update({'source_id': target_node.id})
        WorkspaceItemRelation.query.filter_by(target_id=source_node.id, document_id=doc_id).update({'target_id': target_node.id})
        db.session.delete(source_node)
        db.session.commit()

        return jsonify({
            "message": f"Đã gộp thành công node '{source_node.title}' vào '{target_node.title}'.",
            "source_node_id": source_node.id,
            "target_node_id": target_node.id
        })
        
    except Exception as e:
        db.session.rollback()
        print(f"Lỗi trong execute_merge_nodes: {e}")
        traceback.print_exc()
        return jsonify({"error": "Lỗi server khi thực hiện gộp node."}), 500
    
# =============================================================================
# SECTION 5: CUSTOM TEMPLATE FILTERS
# =============================================================================
@app.template_filter('from_json')
def from_json_filter(value):
    if value is None: return []
    try: return json.loads(value)
    except (json.JSONDecodeError, TypeError): return []

# =============================================================================
# SECTION 6: SCRIPT EXECUTION & DB SETUP
# =============================================================================
def create_db():
     if not os.path.exists(os.path.join(basedir, 'studyvault.db')):
        with app.app_context():
            try: db.create_all(); print("Database created!")
            except Exception as e: print(f"Error creating database: {e}")

def backfill_normalized_names():
    with app.app_context():
        print("Starting to backfill normalized filenames...")
       
        docs = Document.query.filter(Document.filename_normalized.is_(None)).all()
        if not docs:
            print("No documents to update. All filenames are already normalized.")
            return

        for doc in docs:
            doc.filename_normalized = normalize_vietnamese(doc.filename)
        try:
            db.session.commit()
            print(f"Successfully updated {len(docs)} documents.")
        except Exception as e:
            db.session.rollback()
            print(f"An error occurred during backfill: {e}")

if __name__ == '__main__':
    create_db()
 
    with app.app_context(): 
        backfill_normalized_names()
    app.run(debug=True)